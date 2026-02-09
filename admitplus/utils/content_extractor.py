import logging
from typing import Dict, Any
from io import BytesIO

from admitplus.common.exceptions import ContentExtractionError

try:
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PyPDF2 not available, PDF extraction will be disabled")

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available, DOCX extraction will be disabled")


class FileContentExtractor:
    """
    File content extractor - supports PDF and DOCX formats.
    """

    def __init__(self):
        """
        Initialize the file content extractor.
        """
        self.supported_formats = {
            "pdf": PDF_AVAILABLE,
            "docx": DOCX_AVAILABLE,
            "doc": False,
        }

    def extract_text(
        self, file_content: bytes, file_name: str, content_type: str
    ) -> Dict[str, Any]:
        """
        Extract text from file content.

        """
        logging.info(
            f"[Content Extractor] [Extract Text] Starting extraction for file: {file_name}, content-type: {content_type}"
        )

        try:
            file_ext = self._get_file_extension(file_name, content_type)
            logging.info(
                f"[Content Extractor] [Extract Text] Detected file extension: {file_ext}"
            )

            if file_ext == "pdf" and PDF_AVAILABLE:
                return self._extract_pdf_text(file_content, file_name)
            elif file_ext == "docx" and DOCX_AVAILABLE:
                return self._extract_docx_text(file_content, file_name)
            else:
                raise ContentExtractionError(
                    f"Unsupported file format: {file_ext}. Only PDF and DOCX are supported on this system."
                )

        except ContentExtractionError:
            raise
        except Exception as e:
            logging.error(
                f"[Content Extractor] [Extract Text] Extraction failed for {file_name}: {str(e)}"
            )
            raise ContentExtractionError(f"Text extraction failed: {str(e)}")

    def extract_pdf_text_simple(self, file_content: bytes) -> str:
        """
        Simple extraction of text content from PDF files.

        """
        if not PDF_AVAILABLE:
            raise ContentExtractionError("PDF extraction requires PyPDF2 library")

        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            if pdf_reader.is_encrypted:
                if not pdf_reader.decrypt(""):
                    raise ContentExtractionError(
                        "PDF is encrypted and cannot be decrypted"
                    )

            extracted_text = ""

            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    extracted_text += page_text + "\n\n"

            return extracted_text.strip()

        except Exception as e:
            logging.error(
                f"[Content Extractor] [Extract PDF Simple] Failed to extract PDF: {str(e)}"
            )
            raise ContentExtractionError(f"PDF extraction failed: {str(e)}")

    def extract_docx_text_simple(self, file_content: bytes) -> str:
        """
        Simple extraction of text content from DOCX files.

        """
        if not DOCX_AVAILABLE:
            raise ContentExtractionError("DOCX extraction requires python-docx library")

        try:
            docx_file = BytesIO(file_content)
            doc = Document(docx_file)

            extracted_text = ""

            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    extracted_text += para_text + "\n\n"

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            extracted_text += cell_text + "\t"
                    extracted_text += "\n"
                extracted_text += "\n"

            return extracted_text.strip()

        except Exception as e:
            logging.error(
                f"[Content Extractor] [Extract DOCX Simple] Failed to extract DOCX: {str(e)}"
            )
            raise ContentExtractionError(f"DOCX extraction failed: {str(e)}")

    def _get_file_extension(self, file_name: str, content_type: str) -> str:
        """
        Get file extension based on file name and content-type.

        """
        if "." in file_name:
            ext = file_name.lower().split(".")[-1]
            if ext in ["pdf", "docx"]:
                return ext

        content_type = content_type.lower()
        if "pdf" in content_type:
            return "pdf"
        elif (
            "vnd.openxmlformats-officedocument.wordprocessingml.document"
            in content_type
        ):
            return "docx"
        elif "msword" in content_type:
            raise ContentExtractionError(
                "DOC format is not supported on Linux systems. Please convert to PDF or DOCX."
            )
        else:
            if file_name.lower().endswith(".pdf"):
                return "pdf"
            elif file_name.lower().endswith(".docx"):
                return "docx"
            elif file_name.lower().endswith(".doc"):
                raise ContentExtractionError(
                    "DOC format is not supported on Linux systems. Please convert to PDF or DOCX."
                )
            else:
                raise ContentExtractionError(
                    f"Cannot determine file format from filename: {file_name} and content-type: {content_type}"
                )

    def _extract_pdf_text(self, file_content: bytes, file_name: str) -> Dict[str, Any]:
        """
        Extract text content from PDF file.

        """
        logging.info(
            f"[Content Extractor] [Extract PDF] Extracting text from PDF: {file_name}"
        )

        if not PDF_AVAILABLE:
            raise ContentExtractionError("PDF extraction requires PyPDF2 library")

        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            if pdf_reader.is_encrypted:
                logging.warning(
                    f"[Content Extractor] [Extract PDF] PDF is encrypted, attempting to decrypt: {file_name}"
                )
                if not pdf_reader.decrypt(""):
                    raise ContentExtractionError(
                        "PDF is encrypted and cannot be decrypted"
                    )

            total_pages = len(pdf_reader.pages)
            extracted_text = ""
            page_texts = []

            logging.info(
                f"[Content Extractor] [Extract PDF] Processing {total_pages} pages: {file_name}"
            )

            for page_num in range(total_pages):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()

                    if page_text and page_text.strip():
                        extracted_text += page_text + "\n\n"
                        page_texts.append(
                            {
                                "page_number": page_num + 1,
                                "text": page_text.strip(),
                                "char_count": len(page_text),
                            }
                        )
                    else:
                        page_texts.append(
                            {"page_number": page_num + 1, "text": "", "char_count": 0}
                        )

                except Exception as page_error:
                    logging.warning(
                        f"[Content Extractor] [Extract PDF] Error extracting page {page_num + 1}: {str(page_error)}"
                    )
                    page_texts.append(
                        {
                            "page_number": page_num + 1,
                            "text": "",
                            "char_count": 0,
                            "error": str(page_error),
                        }
                    )

            total_chars = len(extracted_text)
            non_empty_pages = sum(1 for page in page_texts if page["char_count"] > 0)

            logging.info(
                f"[Content Extractor] [Extract PDF] Successfully extracted PDF: {file_name}, "
                f"Pages: {total_pages}, Non-empty: {non_empty_pages}, Total chars: {total_chars}"
            )

            result = {
                "success": True,
                "file_type": "pdf",
                "file_name": file_name,
                "total_pages": total_pages,
                "non_empty_pages": non_empty_pages,
                "total_characters": total_chars,
                "extracted_text": extracted_text.strip(),
                "page_details": page_texts,
            }

            try:
                if hasattr(pdf_reader, "metadata") and pdf_reader.metadata:

                    def safe_get_date(metadata, attr_name):
                        """
                        Safely get date field to avoid parsing errors.

                        """
                        try:
                            date_value = getattr(metadata, attr_name, None)
                            if date_value is None:
                                return ""
                            if isinstance(date_value, str):
                                return date_value
                            if hasattr(date_value, "strftime"):
                                return date_value.strftime("%Y-%m-%d %H:%M:%S")
                            return str(date_value)
                        except Exception as e:
                            logging.warning(
                                f"[Content Extractor] [Extract PDF] Failed to extract {attr_name}: {str(e)}"
                            )
                            return ""

                    def safe_get_attr(metadata, attr_name, default=""):
                        """
                        Safely get metadata attribute.

                        """
                        try:
                            return getattr(metadata, attr_name, default) or default
                        except Exception as e:
                            logging.warning(
                                f"[Content Extractor] [Extract PDF] Failed to extract {attr_name}: {str(e)}"
                            )
                            return default

                    result["metadata"] = {
                        "author": safe_get_attr(pdf_reader.metadata, "author"),
                        "title": safe_get_attr(pdf_reader.metadata, "title"),
                        "subject": safe_get_attr(pdf_reader.metadata, "subject"),
                        "creator": safe_get_attr(pdf_reader.metadata, "creator"),
                        "producer": safe_get_attr(pdf_reader.metadata, "producer"),
                        "creation_date": safe_get_date(
                            pdf_reader.metadata, "creation_date"
                        ),
                        "modification_date": safe_get_date(
                            pdf_reader.metadata, "modification_date"
                        ),
                    }
            except Exception as metadata_error:
                logging.warning(
                    f"[Content Extractor] [Extract PDF] Failed to extract metadata: {str(metadata_error)}"
                )
                pass

            return result

        except Exception as e:
            logging.error(
                f"[Content Extractor] [Extract PDF] Failed to extract PDF {file_name}: {str(e)}"
            )
            raise ContentExtractionError(f"PDF extraction failed: {str(e)}")

    def _extract_docx_text(self, file_content: bytes, file_name: str) -> Dict[str, Any]:
        """
        Extract text content from DOCX file.

        """
        logging.info(
            f"[Content Extractor] [Extract DOCX] Extracting text from DOCX: {file_name}"
        )

        if not DOCX_AVAILABLE:
            raise ContentExtractionError("DOCX extraction requires python-docx library")

        try:
            docx_file = BytesIO(file_content)
            doc = Document(docx_file)

            extracted_text = ""
            paragraph_details = []
            total_paragraphs = 0

            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                para_text = paragraph.text.strip()
                if para_text:
                    extracted_text += para_text + "\n\n"
                    paragraph_details.append(
                        {
                            "paragraph_number": para_num,
                            "text": para_text,
                            "char_count": len(para_text),
                            "style": paragraph.style.name
                            if paragraph.style
                            else "Normal",
                        }
                    )
                total_paragraphs += 1

            table_details = []
            for table_num, table in enumerate(doc.tables, 1):
                table_text = ""
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            table_text += cell_text + "\t"
                    table_text += "\n"

                if table_text.strip():
                    extracted_text += f"Table {table_num}:\n{table_text}\n\n"
                    table_details.append(
                        {
                            "table_number": table_num,
                            "text": table_text.strip(),
                            "char_count": len(table_text),
                        }
                    )

            total_chars = len(extracted_text)
            non_empty_paragraphs = sum(
                1 for p in paragraph_details if p["char_count"] > 0
            )

            logging.info(
                f"[Content Extractor] [Extract DOCX] Successfully extracted DOCX: {file_name}, "
                f"Paragraphs: {total_paragraphs}, Non-empty: {non_empty_paragraphs}, "
                f"Tables: {len(table_details)}, Total chars: {total_chars}"
            )

            result = {
                "success": True,
                "file_type": "docx",
                "file_name": file_name,
                "total_paragraphs": total_paragraphs,
                "non_empty_paragraphs": non_empty_paragraphs,
                "total_tables": len(table_details),
                "total_characters": total_chars,
                "extracted_text": extracted_text.strip(),
                "paragraph_details": paragraph_details,
                "table_details": table_details,
            }

            if hasattr(doc, "core_properties"):
                result["core_properties"] = {
                    "title": getattr(doc.core_properties, "title", ""),
                    "subject": getattr(doc.core_properties, "subject", ""),
                    "author": getattr(doc.core_properties, "author", ""),
                    "keywords": getattr(doc.core_properties, "keywords", ""),
                    "comments": getattr(doc.core_properties, "comments", ""),
                    "created": getattr(doc.core_properties, "created", ""),
                    "modified": getattr(doc.core_properties, "modified", ""),
                }

            return result

        except Exception as e:
            logging.error(
                f"[Content Extractor] [Extract DOCX] Failed to extract DOCX {file_name}: {str(e)}"
            )
            raise ContentExtractionError(f"DOCX extraction failed: {str(e)}")

    def is_format_supported(self, file_name: str, content_type: str) -> bool:
        """
        Check if file format is supported.

        """
        try:
            file_ext = self._get_file_extension(file_name, content_type)
            return self.supported_formats.get(file_ext, False)
        except ContentExtractionError:
            return False

    def get_supported_formats(self) -> Dict[str, bool]:
        """
        Get list of supported formats.

        """
        return self.supported_formats.copy()

    def get_system_info(self) -> Dict[str, Any]:
        """
        Get system support information.

        """
        return {
            "supported_formats": self.supported_formats,
            "platform": "linux",
            "pdf_support": PDF_AVAILABLE,
            "docx_support": DOCX_AVAILABLE,
            "doc_support": False,
            "notes": "DOC format is not supported on Linux systems",
        }


"""
Create global instance.

"""
content_extractor = FileContentExtractor()
